import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import pandas as pd
import subprocess
import datetime
from openpyxl import load_workbook
from openpyxl.styles import Font
from docx import Document
from docx.shared import Inches


class StockAnalysisApp:
    def __init__(self, root):
        self.root = root
        self.root.title("收发存分析2.0 By Aaron 2026.3.17")
        self.root.geometry("1400x900")
        self.root.resizable(True, True)
        
        self.file_path = None
        self.df = None
        self.filtered_df = None
        # 保存当前显示的 DataFrame（用于排序时保持列筛选）
        self.current_display_df = None
        
        self.location_options = ["BOS", "HUO", "SFO", "SAN", "LAX", "其他"]
        self.location_vars = {}
        self.group_by_category = tk.BooleanVar(value=True)
        
        # 全选/全不选状态
        self.select_all_state = True  # True 表示当前是全选状态
        # 全部展开/全部折叠状态
        self.expand_all_state = True  # True 表示当前是展开状态
        # 当前视图类型（None, opening, in, out, ending）
        self.current_view_type = None
        
        self.create_widgets()
    
    def create_widgets(self):
        # 顶部文件选择区域
        file_frame = tk.Frame(self.root)
        file_frame.pack(pady=5, fill=tk.X, padx=10)
        
        tk.Label(file_frame, text="文件路径:").pack(side=tk.LEFT, padx=3)
        
        self.path_entry = tk.Entry(file_frame, width=40)
        self.path_entry.pack(side=tk.LEFT, padx=3, fill=tk.X, expand=True)
        
        browse_button = tk.Button(file_frame, text="选择并读取", command=self.select_and_read_file, bg="#4A90E2", fg="white", width=7)
        browse_button.pack(side=tk.LEFT, padx=3)
        
        self.overview_button = tk.Button(file_frame, text="整体概览", command=self.show_overview, bg="#4CAF50", fg="white", width=8)
        self.overview_button.pack(side=tk.LEFT, padx=3)
        
        # 紧凑的控制区域
        control_frame = tk.Frame(self.root)
        control_frame.pack(pady=5, fill=tk.X, padx=10)
        
        # 左侧按钮区域
        button_frame = tk.Frame(control_frame)
        button_frame.pack(side=tk.LEFT, fill=tk.Y)
        
        # 功能按钮 - 紧凑排列
        function_buttons = [
            ("月初", lambda: self.filter_by_type("opening")),
            ("采购", lambda: self.filter_by_type("in")),
            ("耗用", lambda: self.filter_by_type("out")),
            ("结存", lambda: self.filter_by_type("ending")),
            ("excel", self.export_reconciliation),
            ("word", self.export_word_report_v2),
            ("目录", self.open_export_folder)
        ]
        
        for text, cmd in function_buttons:
            btn = tk.Button(button_frame, text=text, command=cmd, width=5, bg="#4A90E2", fg="white")
            btn.pack(side=tk.LEFT, padx=2, pady=1)
        
        # 中间金额汇总区域
        info_frame = tk.Frame(control_frame)
        info_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5)
        
        # 金额标签 - 紧凑横向排列
        amounts = [
            ("期初:", "opening_amount_var"),
            ("采购:", "in_amount_var"),
            ("耗用:", "out_amount_var"),
            ("结存:", "ending_amount_var")
        ]
        
        for label_text, var_name in amounts:
            label = tk.Label(info_frame, text=label_text, font=('Arial', 9, 'bold'))
            label.pack(side=tk.LEFT, padx=3, pady=1)
            var = tk.StringVar(value="$0.00")
            setattr(self, var_name, var)
            value_label = tk.Label(info_frame, textvariable=var, font=('Arial', 9))
            value_label.pack(side=tk.LEFT, padx=3, pady=1)
        
        # 右侧Location筛选区域 - 所有元素一行排列，靠左
        filter_frame = tk.Frame(control_frame)
        filter_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Location复选框 - 单行紧凑排列
        for location in self.location_options:
            var = tk.BooleanVar(value=True)
            self.location_vars[location] = var
            cb = tk.Checkbutton(filter_frame, text=location, variable=var, command=self.apply_location_filter)
            cb.pack(side=tk.LEFT, padx=2, pady=1)
        
        # 控制按钮 - 紧跟在复选框后面
        self.select_all_button = tk.Button(filter_frame, text="全不选", command=self.toggle_select_all, width=5, bg="#4A90E2", fg="white")
        self.select_all_button.pack(side=tk.LEFT, padx=2, pady=1)
        
        self.expand_all_button = tk.Button(filter_frame, text="折叠", command=self.toggle_expand_all, width=5, bg="#4A90E2", fg="white")
        self.expand_all_button.pack(side=tk.LEFT, padx=2, pady=1)
        
        # 分组复选框
        group_checkbox = tk.Checkbutton(filter_frame, text="分组", variable=self.group_by_category, command=self.refresh_display)
        group_checkbox.pack(side=tk.LEFT, padx=2, pady=1)
        
        # 最大化的表格区域
        table_frame = tk.Frame(self.root)
        table_frame.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)
        
        self.tree_scroll_y = ttk.Scrollbar(table_frame, orient=tk.VERTICAL)
        self.tree_scroll_x = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL)
        
        self.tree = ttk.Treeview(
            table_frame,
            yscrollcommand=self.tree_scroll_y.set,
            xscrollcommand=self.tree_scroll_x.set,
            show="tree headings"
        )
        
        self.tree_scroll_y.config(command=self.tree.yview)
        self.tree_scroll_x.config(command=self.tree.xview)
        
        self.tree_scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree_scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 底部状态栏
        status_frame = tk.Frame(self.root)
        status_frame.pack(pady=2, fill=tk.X, padx=10)
        
        self.status_label = tk.Label(status_frame, text="就绪", fg="gray", font=('Arial', 9))
        self.status_label.pack(side=tk.LEFT)
        
        # 信息显示标签
        self.info_label = tk.Label(status_frame, text="", fg="blue", font=('Arial', 9))
        self.info_label.pack(side=tk.LEFT, padx=10)
    
    def select_and_read_file(self):
        file_types = [
            ("Excel文件", "*.xlsx *.xls"),
            ("所有文件", "*.*")
        ]
        
        file_path = filedialog.askopenfilename(
            title="选择收发存汇总表",
            filetypes=file_types
        )
        
        if file_path:
            self.file_path = file_path
            self.path_entry.delete(0, tk.END)
            self.path_entry.insert(0, file_path)
            self.read_excel(file_path)
        else:
            self.status_label.config(text="未选择文件", fg="gray")
    
    def clear_table(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
    
    def refresh_display(self):
        self.display_data_in_table()
    
    def calculate_total_row(self, df, columns, label="合计"):
        numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns
        total_row = []
        for i, col in enumerate(columns):
            if i == 0:
                total_row.append(label)
            elif col in numeric_cols:
                total_value = df[col].sum()
                if isinstance(total_value, float):
                    total_text = f"{total_value:.2f}"
                else:
                    total_text = str(total_value)
                total_row.append(total_text)
            else:
                total_row.append("")
        return total_row
    
    def treeview_sort_column(self, col, reverse):
        # 使用当前显示的 DataFrame，而不是 filtered_df
        # 这样可以保持采购/耗用/结存等视图的列筛选
        display_df = self.current_display_df if self.current_display_df is not None else (self.filtered_df if self.filtered_df is not None else self.df)
        if display_df is None or len(display_df) == 0:
            return
        
        # 创建副本，避免修改原始数据
        display_df = display_df.copy()
        
        # 尝试将列转换为数值类型进行排序
        try:
            # 检查列是否可以转换为数值
            display_df[col] = pd.to_numeric(display_df[col], errors='coerce')
            # 按数值排序
            display_df = display_df.sort_values(by=col, ascending=not reverse)
        except:
            # 如果转换失败，按字符串排序
            display_df = display_df.sort_values(by=col, ascending=not reverse)
        
        # 重新显示数据（保持当前的列筛选）
        self.display_data_in_table(display_df)
        
        # 更新表头排序指示（只更新当前显示的列）
        for c in self.tree["columns"]:
            if c == col:
                # 设置排序指示
                if reverse:
                    self.tree.heading(c, text=f"{c} ↑", command=lambda _col=col: self.treeview_sort_column(_col, False))
                else:
                    self.tree.heading(c, text=f"{c} ↓", command=lambda _col=col: self.treeview_sort_column(_col, True))
            else:
                # 重置其他列的排序指示
                self.tree.heading(c, text=str(c), command=lambda _col=c: self.treeview_sort_column(_col, False))
    
    def display_data_in_table(self, display_df=None):
        self.clear_table()
        
        if display_df is None:
            display_df = self.filtered_df if self.filtered_df is not None else self.df
        
        if display_df is None or len(display_df) == 0:
            # 即使没有数据，也要清除 current_display_df
            self.current_display_df = None
            return
        
        # 保存当前显示的 DataFrame（包括列筛选信息）
        # 这对于排序时保持采购/耗用/结存/月初视图的列筛选至关重要
        self.current_display_df = display_df.copy()
        
        columns = list(display_df.columns)
        self.tree["columns"] = columns
        
        for col in columns:
            # 添加排序功能
            self.tree.heading(col, text=str(col), command=lambda _col=col: self.treeview_sort_column(_col, False))
            max_width = max(len(str(col)), display_df[col].astype(str).str.len().max() if len(display_df) > 0 else 10)
            col_width = min(max_width * 10 + 20, 200)
            self.tree.column(col, width=col_width)
        
        self.tree.column("#0", width=20, minwidth=20, stretch=tk.NO)
        
        if not self.group_by_category.get():
            for i, row in display_df.iterrows():
                self.tree.insert("", tk.END, values=list(row))
            
            total_row = self.calculate_total_row(display_df, columns)
            total_item = self.tree.insert("", tk.END, values=total_row)
            self.tree.item(total_item, tags=('total',))
            self.tree.tag_configure('total', background='#E0E0E0', font=('Arial', 9, 'bold'))
        else:
            category_col = None
            # 优先查找中文列名
            for col in columns:
                col_str = str(col)
                if "类别" in col_str or "category" in col_str.lower():
                    category_col = col
                    break
            
            if category_col is None:
                for i, row in display_df.iterrows():
                    self.tree.insert("", tk.END, values=list(row))
                
                total_row = self.calculate_total_row(display_df, columns)
                total_item = self.tree.insert("", tk.END, values=total_row)
                self.tree.item(total_item, tags=('total',))
                self.tree.tag_configure('total', background='#E0E0E0', font=('Arial', 9, 'bold'))
                return
            
            categories = display_df[category_col].unique()
            
            for category in categories:
                category_df = display_df[display_df[category_col] == category]
                
                category_total = self.calculate_total_row(category_df, columns, label=f"{category} ({len(category_df)}条)")
                category_item = self.tree.insert("", tk.END, values=category_total)
                self.tree.item(category_item, tags=('category',))
                self.tree.tag_configure('category', background='#C8E6C9', font=('Arial', 9, 'bold'))
                
                for i, row in category_df.iterrows():
                    self.tree.insert(category_item, tk.END, values=list(row))
                
                self.tree.item(category_item, open=True)
            
            total_row = self.calculate_total_row(display_df, columns)
            total_item = self.tree.insert("", tk.END, values=total_row)
            self.tree.item(total_item, tags=('total',))
            self.tree.tag_configure('total', background='#E0E0E0', font=('Arial', 9, 'bold'))
    
    def calculate_amounts(self):
        if self.df is None:
            return
        
        try:
            df = self.filtered_df if self.filtered_df is not None else self.df
            
            # 计算期初金额
            opening_cols = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col)]
            opening_amount = 0
            if opening_cols:
                numeric_df = df[opening_cols].apply(pd.to_numeric, errors='coerce')
                opening_amount = numeric_df.sum().sum()
            
            # 计算采购金额
            in_cols = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
            in_amount = 0
            if in_cols:
                numeric_df = df[in_cols].apply(pd.to_numeric, errors='coerce')
                in_amount = numeric_df.sum().sum()
            
            # 计算耗用金额
            out_cols = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col)]
            out_amount = 0
            if out_cols:
                numeric_df = df[out_cols].apply(pd.to_numeric, errors='coerce')
                out_amount = numeric_df.sum().sum()
            
            # 计算结存金额
            ending_cols = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col)]
            ending_amount = 0
            if ending_cols:
                numeric_df = df[ending_cols].apply(pd.to_numeric, errors='coerce')
                ending_amount = numeric_df.sum().sum()
            
            # 更新UI显示
            self.opening_amount_var.set(f"${opening_amount:,.2f}")
            self.in_amount_var.set(f"${in_amount:,.2f}")
            self.out_amount_var.set(f"${out_amount:,.2f}")
            self.ending_amount_var.set(f"${ending_amount:,.2f}")
            
        except Exception as e:
            print(f"计算金额失败: {str(e)}")
    
    def apply_location_filter(self):
        if self.df is None:
            return
        
        try:
            selected_locations = [loc for loc, var in self.location_vars.items() if var.get()]
            
            if not selected_locations:
                self.filtered_df = self.df.head(0).copy()
                self.info_label.config(text=f"文件: {os.path.basename(self.file_path)} | 原始行数: {len(self.df)} | 筛选后行数: 0")
                self.display_data_in_table()
                self.status_label.config(text="无筛选结果", fg="orange")
                # 更新金额汇总
                self.calculate_amounts()
                return
            
            location_col = None
            # 优先查找中文列名
            for col in self.df.columns:
                col_str = str(col)
                if "仓库" in col_str or "location" in col_str.lower():
                    location_col = col
                    break
            
            if location_col is None:
                location_col = self.df.columns[0]
            
            other_selected = "其他" in selected_locations
            other_locations = [loc for loc in selected_locations if loc != "其他"]
            
            main_locations = ["HUO", "SAN", "BOS", "SFO", "LAX"]
            
            if other_selected and other_locations:
                # 包含选择的仓库和其他仓库
                mask_other = ~self.df[location_col].astype(str).str.contains('|'.join(main_locations), case=False, na=False)
                mask_selected = self.df[location_col].astype(str).str.contains('|'.join(other_locations), case=False, na=False)
                mask = mask_other | mask_selected
                self.filtered_df = self.df[mask].copy()
            elif other_selected:
                # 只包含其他仓库（不包含HUO, SAN, BOS, SFO, LAX的记录）
                mask_other = ~self.df[location_col].astype(str).str.contains('|'.join(main_locations), case=False, na=False)
                self.filtered_df = self.df[mask_other].copy()
            else:
                # 只包含选择的仓库
                mask = self.df[location_col].astype(str).str.contains('|'.join(other_locations), case=False, na=False)
                self.filtered_df = self.df[mask].copy()
            
            self.info_label.config(text=f"文件: {os.path.basename(self.file_path)} | 原始行数: {len(self.df)} | 筛选后行数: {len(self.filtered_df)}")
            self.display_data_in_table()
            
            # 重置视图类型和当前显示数据
            self.current_view_type = None
            self.current_display_df = None
            
            # 更新金额汇总
            self.calculate_amounts()
            
            self.status_label.config(text=f"筛选完成，显示 {len(self.filtered_df)} 条记录", fg="green")
            
        except Exception as e:
            messagebox.showerror("错误", f"筛选失败: {str(e)}")
            self.status_label.config(text="筛选失败", fg="red")
    
    def select_all(self):
        for loc in self.location_options:
            self.location_vars[loc].set(True)
        self.apply_location_filter()
        self.status_label.config(text="已全选", fg="green")
        # 更新按钮状态
        self.select_all_button.config(text="全不选", bg="#F44336")
        self.select_all_state = True
    
    def deselect_all(self):
        for loc in self.location_options:
            self.location_vars[loc].set(False)
        self.apply_location_filter()
        self.status_label.config(text="已全不选", fg="green")
        # 更新按钮状态
        self.select_all_button.config(text="全选", bg="#4CAF50")
        self.select_all_state = False
    
    def toggle_select_all(self):
        if self.select_all_state:
            # 当前是全选状态，切换到全不选
            self.deselect_all()
            self.select_all_button.config(text="全选", bg="#4CAF50")
            self.select_all_state = False
        else:
            # 当前是全不选状态，切换到全选
            self.select_all()
            self.select_all_button.config(text="全不选", bg="#F44336")
            self.select_all_state = True
    
    def filter_by_type(self, type_filter):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取Excel文件")
            return
        
        try:
            df = self.filtered_df if self.filtered_df is not None else self.df
            
            type_filter = type_filter.lower()
            
            if type_filter == "opening":
                # 查找期初相关列
                opening_cols = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col)]
                if opening_cols:
                    mask = (df[opening_cols] != 0).any(axis=1)
                    filtered_df = df[mask].copy()
                    # 显示相关列
                    display_cols = [col for col in filtered_df.columns if "opening" in str(col).lower() or "期初" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                    display_df = filtered_df[display_cols]
                    self.info_label.config(text=f"本月初存货 | 显示 {len(display_df)} 条记录")
                    self.display_data_in_table(display_df)
                    # 更新金额汇总
                    self.calculate_amounts()
                    self.status_label.config(text="本月初存货筛选完成", fg="green")
                    # 设置当前视图类型
                    self.current_view_type = "opening"
                else:
                    messagebox.showwarning("警告", "未找到期初相关列")
            
            elif type_filter == "in":
                # 查找入库相关列
                in_cols = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
                if in_cols:
                    # 确保只显示至少有一个入库字段不为 0 的行
                    numeric_df = df[in_cols].apply(pd.to_numeric, errors='coerce')
                    mask = (numeric_df != 0).any(axis=1)
                    filtered_df = df[mask].copy()
                    # 显示相关列
                    display_cols = [col for col in filtered_df.columns if ("in" in str(col).lower() or "入库" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()) and not ("ending" in str(col).lower() or "期末" in str(col) or "opening" in str(col).lower() or "期初" in str(col))]
                    display_df = filtered_df[display_cols]
                    self.info_label.config(text=f"本月采购 | 显示 {len(display_df)} 条记录")
                    self.display_data_in_table(display_df)
                    # 更新金额汇总
                    self.calculate_amounts()
                    self.status_label.config(text="本月采购筛选完成", fg="green")
                    # 设置当前视图类型
                    self.current_view_type = "in"
                else:
                    messagebox.showwarning("警告", "未找到入库相关列")
            
            elif type_filter == "out":
                # 查找出库相关列
                out_cols = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col)]
                if out_cols:
                    mask = (df[out_cols] != 0).any(axis=1)
                    filtered_df = df[mask].copy()
                    # 显示相关列
                    display_cols = [col for col in filtered_df.columns if "out" in str(col).lower() or "出库" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                    display_df = filtered_df[display_cols]
                    self.info_label.config(text=f"本月耗用 | 显示 {len(display_df)} 条记录")
                    self.display_data_in_table(display_df)
                    # 更新金额汇总
                    self.calculate_amounts()
                    self.status_label.config(text="本月耗用筛选完成", fg="green")
                    # 设置当前视图类型
                    self.current_view_type = "out"
                else:
                    messagebox.showwarning("警告", "未找到出库相关列")
            
            elif type_filter == "ending":
                # 查找期末相关列
                ending_cols = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col)]
                if ending_cols:
                    mask = (df[ending_cols] != 0).any(axis=1)
                    filtered_df = df[mask].copy()
                    # 显示相关列
                    display_cols = [col for col in filtered_df.columns if "ending" in str(col).lower() or "期末" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                    display_df = filtered_df[display_cols]
                    self.info_label.config(text=f"月底留存 | 显示 {len(display_df)} 条记录")
                    self.display_data_in_table(display_df)
                    # 更新金额汇总
                    self.calculate_amounts()
                    self.status_label.config(text="月底留存筛选完成", fg="green")
                    # 设置当前视图类型
                    self.current_view_type = "ending"
                else:
                    messagebox.showwarning("警告", "未找到期末相关列")
                    
        except Exception as e:
            messagebox.showerror("错误", f"筛选失败: {str(e)}")
            self.status_label.config(text="筛选失败", fg="red")
    
    def reset_type_filter(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取 Excel 文件")
            return
        
        # 重置视图类型和当前显示数据
        self.current_view_type = None
        self.current_display_df = None
        
        # 重置为 location 筛选后的数据
        self.display_data_in_table()
        # 更新 info_label
        if self.filtered_df is not None:
            self.info_label.config(text=f"文件：{os.path.basename(self.file_path)} | 原始行数：{len(self.df)} | 筛选后行数：{len(self.filtered_df)}")
        else:
            self.info_label.config(text=f"文件：{os.path.basename(self.file_path)} | 行数：{len(self.df)} | 列数：{len(self.df.columns)}")
        # 更新金额汇总
        self.calculate_amounts()
        self.status_label.config(text="类型筛选已重置", fg="green")
    
    def export_reconciliation(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取Excel文件")
            return
        
        try:
            self.status_label.config(text="正在导出对账数据...", fg="orange")
            self.root.update()
            
            # 直接使用 filtered_df（已经应用了 Location 筛选）
            # 不再重复应用筛选逻辑，避免双重筛选
            filtered_df = self.filtered_df if self.filtered_df is not None else self.df
            
            if len(filtered_df) == 0:
                messagebox.showwarning("警告", "没有找到符合条件的数据")
                self.status_label.config(text="导出失败", fg="red")
                return
            
            export_dir = os.path.join(os.path.dirname(self.file_path), "导出") if self.file_path else os.path.join(os.getcwd(), "导出")
            if not os.path.exists(export_dir):
                os.makedirs(export_dir)
            
            # 获取用户选择的 site 名称
            selected_locations = [loc for loc, var in self.location_vars.items() if var.get()]
            site_name = "_&".join(selected_locations) if selected_locations else "all"
            
            # 清理目录内所有类似命名的文件
            for file in os.listdir(export_dir):
                if file.endswith(".xlsx") and "-stock-balance-detail_" in file:
                    file_path = os.path.join(export_dir, file)
                    try:
                        os.remove(file_path)
                    except:
                        pass
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            export_file = os.path.join(export_dir, f"{site_name}-stock-balance-detail_{timestamp}.xlsx")
            
            # 导出为一个Excel文件，包含期初、本期采购、本期耗用、本期结存
            with pd.ExcelWriter(export_file, engine='openpyxl') as writer:
                # 期初数据
                opening_cols = [col for col in filtered_df.columns if "opening" in str(col).lower() or "期初" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                if opening_cols:
                    # 只包含期初字段不为0的行
                    opening_cols_filter = [col for col in filtered_df.columns if "opening" in str(col).lower() or "期初" in str(col)]
                    if opening_cols_filter:
                        numeric_df = filtered_df[opening_cols_filter].apply(pd.to_numeric, errors='coerce')
                        opening_mask = (numeric_df != 0).any(axis=1)
                        opening_data = filtered_df[opening_mask][opening_cols]
                    else:
                        opening_data = filtered_df[opening_cols]
                    opening_data.to_excel(writer, sheet_name="期初", index=False)
                
                # 本期采购数据
                in_cols = [col for col in filtered_df.columns if ("in" in str(col).lower() or "入库" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()) and not ("ending" in str(col).lower() or "期末" in str(col) or "opening" in str(col).lower() or "期初" in str(col))]
                if in_cols:
                    # 只包含入库字段不为0的行
                    in_cols_filter = [col for col in filtered_df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
                    if in_cols_filter:
                        numeric_df = filtered_df[in_cols_filter].apply(pd.to_numeric, errors='coerce')
                        in_mask = (numeric_df != 0).any(axis=1)
                        in_data = filtered_df[in_mask][in_cols]
                    else:
                        in_data = filtered_df[in_cols]
                    in_data.to_excel(writer, sheet_name="本期采购", index=False)
                
                # 本期耗用数据
                out_cols = [col for col in filtered_df.columns if "out" in str(col).lower() or "出库" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                if out_cols:
                    # 只包含出库字段不为0的行
                    out_cols_filter = [col for col in filtered_df.columns if "out" in str(col).lower() or "出库" in str(col)]
                    if out_cols_filter:
                        numeric_df = filtered_df[out_cols_filter].apply(pd.to_numeric, errors='coerce')
                        out_mask = (numeric_df != 0).any(axis=1)
                        out_data = filtered_df[out_mask][out_cols]
                    else:
                        out_data = filtered_df[out_cols]
                    out_data.to_excel(writer, sheet_name="本期耗用", index=False)
                
                # 本期结存数据
                ending_cols = [col for col in filtered_df.columns if "ending" in str(col).lower() or "期末" in str(col) or "仓库" in str(col) or "location" in str(col).lower() or "类别" in str(col) or "category" in str(col).lower() or "产品" in str(col) or "product" in str(col).lower()]
                if ending_cols:
                    # 只包含期末字段不为0的行
                    ending_cols_filter = [col for col in filtered_df.columns if "ending" in str(col).lower() or "期末" in str(col)]
                    if ending_cols_filter:
                        numeric_df = filtered_df[ending_cols_filter].apply(pd.to_numeric, errors='coerce')
                        ending_mask = (numeric_df != 0).any(axis=1)
                        ending_data = filtered_df[ending_mask][ending_cols]
                    else:
                        ending_data = filtered_df[ending_cols]
                    ending_data.to_excel(writer, sheet_name="本期结存", index=False)
            
            # 调整Excel文件的列宽、标题样式，并添加合计行
            wb = load_workbook(export_file)
            bold_font = Font(bold=True)
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                # 设置标题行加粗
                for cell in ws[1]:
                    cell.font = bold_font
                
                # 计算并添加合计行
                if ws.max_row > 1:  # 确保有数据行
                    # 确定数值列
                    numeric_columns = []
                    for col in range(1, ws.max_column + 1):
                        # 检查第二行是否为数值
                        cell_value = ws.cell(row=2, column=col).value
                        try:
                            float(cell_value)
                            numeric_columns.append(col)
                        except:
                            pass
                    
                    # 记录当前最大行号（在添加合计行之前）
                    current_max_row = ws.max_row
                    
                    # 添加合计行
                    total_row = current_max_row + 1
                    ws.cell(row=total_row, column=1, value="合计").font = bold_font
                    
                    # 计算数值列的总和
                    for col in numeric_columns:
                        # 使用Excel公式计算总和，确保不包含合计行本身
                        ws.cell(row=total_row, column=col).value = f"=SUM({ws.cell(row=2, column=col).coordinate}:{ws.cell(row=current_max_row, column=col).coordinate})"
                        ws.cell(row=total_row, column=col).font = bold_font
                
                # 自动调整列宽
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)  # 限制最大列宽
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            wb.save(export_file)
            
            self.status_label.config(text="对账导出完成", fg="green")
            messagebox.showinfo("成功", f"对账数据已导出到:\n{export_file}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {str(e)}")
            self.status_label.config(text="导出失败", fg="red")
    
    def open_export_folder(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取Excel文件")
            return
        
        try:
            export_dir = os.path.join(os.path.dirname(self.file_path), "导出") if self.file_path else os.path.join(os.getcwd(), "导出")
            if not os.path.exists(export_dir):
                os.makedirs(export_dir)
            
            os.startfile(export_dir)
            self.status_label.config(text="已打开导出文件夹", fg="green")
            
        except Exception as e:
            messagebox.showerror("错误", f"打开文件夹失败: {str(e)}")
            self.status_label.config(text="打开失败", fg="red")
    

    
    def collapse_all(self):
        for item in self.tree.get_children():
            self.tree.item(item, open=False)
        self.status_label.config(text="全部折叠", fg="blue")
    
    def expand_all(self):
        for item in self.tree.get_children():
            self.tree.item(item, open=True)
        self.status_label.config(text="全部展开", fg="blue")
    
    def toggle_expand_all(self):
        if self.expand_all_state:
            # 当前是展开状态，切换到折叠
            self.collapse_all()
            self.expand_all_button.config(text="全部展开", bg="#009688")
            self.expand_all_state = False
        else:
            # 当前是折叠状态，切换到展开
            self.expand_all()
            self.expand_all_button.config(text="全部折叠", bg="#795548")
            self.expand_all_state = True
    
    def export_word_report(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取Excel文件")
            return
        
        try:
            self.status_label.config(text="正在导出Word报告...", fg="orange")
            self.root.update()
            
            df = self.filtered_df if self.filtered_df is not None else self.df
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading('收发存汇总表分析报告', 0)
            
            # 添加开头部分
            doc.add_paragraph()
            doc.add_paragraph('Dear Site Head,')
            doc.add_paragraph('Below is the analysis of site inventory movement. For detailed information, please refer to the Excel file.')
            doc.add_paragraph()
            
            # 添加数据表格
            doc.add_paragraph()
            doc.add_heading('期初库存', level=1)
            
            # 期初数据
            opening_cols = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col) or "产品" in str(col) or "product" in str(col).lower()]
            # 过滤掉location、category和produce code列
            opening_cols = [col for col in opening_cols if not any(keyword in str(col).lower() for keyword in ["仓库", "location", "类别", "category", "code", "编码"])]
            
            if opening_cols:
                # 只包含期初字段不为0的行
                opening_cols_filter = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col)]
                if opening_cols_filter:
                    numeric_df = df[opening_cols_filter].apply(pd.to_numeric, errors='coerce')
                    opening_mask = (numeric_df != 0).any(axis=1)
                    opening_data = df[opening_mask][opening_cols]
                else:
                    opening_data = df[opening_cols]
                
                # 添加表格
                table = doc.add_table(rows=1, cols=len(opening_data.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(opening_data.columns):
                    hdr_cells[i].text = str(col)
                
                # 设置表格列宽
                for i, col in enumerate(opening_data.columns):
                    if "产品" in str(col) or "product" in str(col).lower():
                        # 产品名称列设置较宽
                        table.columns[i].width = Inches(3.5)
                    else:
                        # 数值列设置较窄
                        table.columns[i].width = Inches(1.2)
                
                # 添加数据行
                for _, row in opening_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        col_name = opening_data.columns[i]
                        if "产品" in str(col_name) or "product" in str(col_name).lower():
                            # 显示完整产品名称
                            row_cells[i].text = str(value)
                        else:
                            # 数值列右对齐
                            row_cells[i].text = str(value)
                            # 设置单元格对齐方式为右对齐
                            row_cells[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                
                # 添加合计行
                if len(opening_data) > 0:
                    total_row = table.add_row().cells
                    total_row[0].text = "合计"
                    for i in range(1, len(opening_data.columns)):
                        col_name = opening_data.columns[i]
                        if any(keyword in str(col_name).lower() for keyword in ["opening", "期初"]):
                            total_value = opening_data[col_name].sum()
                            if isinstance(total_value, float):
                                total_row[i].text = f"{total_value:.2f}"
                            else:
                                total_row[i].text = str(total_value)
                            # 设置合计行数值列右对齐
                            total_row[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                        else:
                            total_row[i].text = ""
            
            # 本期采购
            doc.add_paragraph()
            doc.add_heading('本期采购', level=1)
            
            in_cols = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col) or "产品" in str(col) or "product" in str(col).lower()) and not ("ending" in str(col).lower() or "期末" in str(col) or "opening" in str(col).lower() or "期初" in str(col))]
            # 过滤掉location、category和produce code列
            in_cols = [col for col in in_cols if not any(keyword in str(col).lower() for keyword in ["仓库", "location", "类别", "category", "code", "编码"])]
            
            if in_cols:
                # 只包含入库字段不为0的行
                in_cols_filter = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
                if in_cols_filter:
                    numeric_df = df[in_cols_filter].apply(pd.to_numeric, errors='coerce')
                    in_mask = (numeric_df != 0).any(axis=1)
                    in_data = df[in_mask][in_cols]
                else:
                    in_data = df[in_cols]
                
                # 添加表格
                table = doc.add_table(rows=1, cols=len(in_data.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(in_data.columns):
                    hdr_cells[i].text = str(col)
                
                # 设置表格列宽
                for i, col in enumerate(in_data.columns):
                    if "产品" in str(col) or "product" in str(col).lower():
                        # 产品名称列设置较宽
                        table.columns[i].width = Inches(3.5)
                    else:
                        # 数值列设置较窄
                        table.columns[i].width = Inches(1.2)
                
                # 添加数据行
                for _, row in in_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        col_name = in_data.columns[i]
                        if "产品" in str(col_name) or "product" in str(col_name).lower():
                            # 显示完整产品名称
                            row_cells[i].text = str(value)
                        else:
                            # 数值列右对齐
                            row_cells[i].text = str(value)
                            # 设置单元格对齐方式为右对齐
                            row_cells[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                
                # 添加合计行
                if len(in_data) > 0:
                    total_row = table.add_row().cells
                    total_row[0].text = "合计"
                    for i in range(1, len(in_data.columns)):
                        col_name = in_data.columns[i]
                        if any(keyword in str(col_name).lower() for keyword in ["in", "入库"]):
                            total_value = in_data[col_name].sum()
                            if isinstance(total_value, float):
                                total_row[i].text = f"{total_value:.2f}"
                            else:
                                total_row[i].text = str(total_value)
                            # 设置合计行数值列右对齐
                            total_row[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                        else:
                            total_row[i].text = ""
            
            # 本期耗用
            doc.add_paragraph()
            doc.add_heading('本期耗用', level=1)
            
            out_cols = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col) or "产品" in str(col) or "product" in str(col).lower()]
            # 过滤掉location、category和produce code列
            out_cols = [col for col in out_cols if not any(keyword in str(col).lower() for keyword in ["仓库", "location", "类别", "category", "code", "编码"])]
            
            if out_cols:
                # 只包含出库字段不为0的行
                out_cols_filter = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col)]
                if out_cols_filter:
                    numeric_df = df[out_cols_filter].apply(pd.to_numeric, errors='coerce')
                    out_mask = (numeric_df != 0).any(axis=1)
                    out_data = df[out_mask][out_cols]
                else:
                    out_data = df[out_cols]
                
                # 添加表格
                table = doc.add_table(rows=1, cols=len(out_data.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(out_data.columns):
                    hdr_cells[i].text = str(col)
                
                # 设置表格列宽
                for i, col in enumerate(out_data.columns):
                    if "产品" in str(col) or "product" in str(col).lower():
                        # 产品名称列设置较宽
                        table.columns[i].width = Inches(3.5)
                    else:
                        # 数值列设置较窄
                        table.columns[i].width = Inches(1.2)
                
                # 添加数据行
                for _, row in out_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        col_name = out_data.columns[i]
                        if "产品" in str(col_name) or "product" in str(col_name).lower():
                            # 显示完整产品名称
                            row_cells[i].text = str(value)
                        else:
                            # 数值列右对齐
                            row_cells[i].text = str(value)
                            # 设置单元格对齐方式为右对齐
                            row_cells[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                
                # 添加合计行
                if len(out_data) > 0:
                    total_row = table.add_row().cells
                    total_row[0].text = "合计"
                    for i in range(1, len(out_data.columns)):
                        col_name = out_data.columns[i]
                        if any(keyword in str(col_name).lower() for keyword in ["out", "出库"]):
                            total_value = out_data[col_name].sum()
                            if isinstance(total_value, float):
                                total_row[i].text = f"{total_value:.2f}"
                            else:
                                total_row[i].text = str(total_value)
                            # 设置合计行数值列右对齐
                            total_row[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                        else:
                            total_row[i].text = ""
            
            # 本期结存
            doc.add_paragraph()
            doc.add_heading('本期结存', level=1)
            
            ending_cols = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col) or "产品" in str(col) or "product" in str(col).lower()]
            # 过滤掉location、category和produce code列
            ending_cols = [col for col in ending_cols if not any(keyword in str(col).lower() for keyword in ["仓库", "location", "类别", "category", "code", "编码"])]
            
            if ending_cols:
                # 只包含期末字段不为0的行
                ending_cols_filter = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col)]
                if ending_cols_filter:
                    numeric_df = df[ending_cols_filter].apply(pd.to_numeric, errors='coerce')
                    ending_mask = (numeric_df != 0).any(axis=1)
                    ending_data = df[ending_mask][ending_cols]
                else:
                    ending_data = df[ending_cols]
                
                # 添加表格
                table = doc.add_table(rows=1, cols=len(ending_data.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(ending_data.columns):
                    hdr_cells[i].text = str(col)
                
                # 设置表格列宽
                for i, col in enumerate(ending_data.columns):
                    if "产品" in str(col) or "product" in str(col).lower():
                        # 产品名称列设置较宽
                        table.columns[i].width = Inches(3.5)
                    else:
                        # 数值列设置较窄
                        table.columns[i].width = Inches(1.2)
                
                # 添加数据行
                for _, row in ending_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        col_name = ending_data.columns[i]
                        if "产品" in str(col_name) or "product" in str(col_name).lower():
                            # 显示完整产品名称
                            row_cells[i].text = str(value)
                        else:
                            # 数值列右对齐
                            row_cells[i].text = str(value)
                            # 设置单元格对齐方式为右对齐
                            row_cells[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                
                # 添加合计行
                if len(ending_data) > 0:
                    total_row = table.add_row().cells
                    total_row[0].text = "合计"
                    for i in range(1, len(ending_data.columns)):
                        col_name = ending_data.columns[i]
                        if any(keyword in str(col_name).lower() for keyword in ["ending", "期末"]):
                            total_value = ending_data[col_name].sum()
                            if isinstance(total_value, float):
                                total_row[i].text = f"{total_value:.2f}"
                            else:
                                total_row[i].text = str(total_value)
                            # 设置合计行数值列右对齐
                            total_row[i].paragraphs[0].alignment = 2  # 2 表示右对齐
                        else:
                            total_row[i].text = ""
            
            # 保存Word文档
            export_dir = os.path.join(os.path.dirname(self.file_path), "导出") if self.file_path else os.path.join(os.getcwd(), "导出")
            if not os.path.exists(export_dir):
                os.makedirs(export_dir)
            
            # 获取用户选择的site名称
            selected_locations = [loc for loc, var in self.location_vars.items() if var.get()]
            site_name = "_&".join(selected_locations) if selected_locations else "all"
            
            # 清理目录内所有类似命名的文件
            for file in os.listdir(export_dir):
                if file.endswith(".docx") and "-stock-report_" in file:
                    file_path = os.path.join(export_dir, file)
                    try:
                        os.remove(file_path)
                    except:
                        pass
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            export_file = os.path.join(export_dir, f"{site_name}-stock-report_{timestamp}.docx")
            
            doc.save(export_file)
            
            self.status_label.config(text="Word报告导出完成", fg="green")
            messagebox.showinfo("成功", f"Word报告已导出到:\n{export_file}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{str(e)}")
            self.status_label.config(text="导出失败", fg="red")
    
    def export_word_report_v2(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取 Excel 文件")
            return
        
        try:
            self.status_label.config(text="正在导出 Word 报告...", fg="orange")
            self.root.update()
            
            df = self.filtered_df if self.filtered_df is not None else self.df
            
            # 创建 Word 文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('Monthly Inventory Movement Analysis', 0)
            title.alignment = 1  # 居中
            
            # 获取用户选择的 site 名称
            selected_locations = [loc for loc, var in self.location_vars.items() if var.get()]
            site_name = "_&".join(selected_locations) if selected_locations else "all"
            
            # 添加称呼
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(f'Dear {site_name} Site Head,')
            run.bold = True
            
            doc.add_paragraph()
            doc.add_paragraph('Please find below the monthly inventory movement analysis for your review.')
            
            # 识别列
            opening_cols = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col)]
            in_cols = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
            out_cols = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col)]
            ending_cols = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col)]
            
            category_col = None
            for col in df.columns:
                col_str = str(col)
                if "类别" in col_str or "category" in col_str.lower():
                    category_col = col
                    break
            
            # 1. Summary 部分
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('●   Summary')
            run.bold = True
            
            # 按类别汇总表格
            if category_col:
                categories = sorted(df[category_col].unique())
                
                # 准备汇总数据
                summary_data = []
                for category in categories:
                    category_df = df[df[category_col] == category]
                    
                    opening_amount = 0
                    if opening_cols:
                        numeric_df = category_df[opening_cols].apply(pd.to_numeric, errors='coerce')
                        opening_amount = numeric_df.sum().sum()
                    
                    in_amount = 0
                    if in_cols:
                        numeric_df = category_df[in_cols].apply(pd.to_numeric, errors='coerce')
                        in_amount = numeric_df.sum().sum()
                    
                    out_amount = 0
                    if out_cols:
                        numeric_df = category_df[out_cols].apply(pd.to_numeric, errors='coerce')
                        out_amount = numeric_df.sum().sum()
                    
                    ending_amount = 0
                    if ending_cols:
                        numeric_df = category_df[ending_cols].apply(pd.to_numeric, errors='coerce')
                        ending_amount = numeric_df.sum().sum()
                    
                    summary_data.append({
                        'category': category,
                        'opening': opening_amount,
                        'in': in_amount,
                        'out': out_amount,
                        'ending': ending_amount
                    })
                
                # 添加合计行
                total_opening = sum(item['opening'] for item in summary_data)
                total_in = sum(item['in'] for item in summary_data)
                total_out = sum(item['out'] for item in summary_data)
                total_ending = sum(item['ending'] for item in summary_data)
                summary_data.append({
                    'category': 'Total',
                    'opening': total_opening,
                    'in': total_in,
                    'out': total_out,
                    'ending': total_ending
                })
                
                # 创建表格 - 使用英文固定表头
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ['Category', 'Opening', 'Purchase', 'Consumption', 'Ending']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                
                # 设置列宽：第一列 50%，其余列平均分配剩余 50%
                table.columns[0].width = Inches(3.0)  # 第一列占 50%
                for i in range(1, 5):
                    table.columns[i].width = Inches(0.75)  # 其余列各占约 12.5%
                
                for item in summary_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item['category'])
                    row_cells[1].text = f"{item['opening']:,.2f}"
                    row_cells[2].text = f"{item['in']:,.2f}"
                    row_cells[3].text = f"{item['out']:,.2f}"
                    row_cells[4].text = f"{item['ending']:,.2f}"
            
            # 2. Inventory Receipt Analysis 部分
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('●   Inventory Receipt Analysis')
            run.bold = True
            
            # 计算当月入库金额和数量
            total_receipt_amount = 0
            total_receipt_qty = 0
            
            if in_cols:
                numeric_df = df[in_cols].apply(pd.to_numeric, errors='coerce')
                total_receipt_amount = numeric_df.sum().sum()
                qty_cols = [col for col in in_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if qty_cols:
                    numeric_qty_df = df[qty_cols].apply(pd.to_numeric, errors='coerce')
                    total_receipt_qty = numeric_qty_df.sum().sum()
                else:
                    total_receipt_qty = len(df[df[in_cols[0]] != 0])
            
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Total inventory receipt amount for the month: ')
            run = para.add_run(f'${total_receipt_amount:,.2f}')
            run.bold = True
            run = para.add_run(', this covers a total of ')
            run = para.add_run(f'{total_receipt_qty:,.0f}')
            run.bold = True
            run = para.add_run(' items, with the top 5 items by value as follows:')
            
            # 获取前 5 个入库物品
            product_col = None
            for col in df.columns:
                col_str = str(col)
                if "产品" in col_str or "product" in col_str.lower():
                    product_col = col
                    break
            
            if product_col and in_cols:
                in_cols_no_qty = [col for col in in_cols if not ("数量" in str(col) or "qty" in str(col).lower())]
                qty_cols_in = [col for col in in_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if in_cols_no_qty:
                    df_with_products = df.copy()
                    df_with_products['total_in'] = df_with_products[in_cols_no_qty].apply(pd.to_numeric, errors='coerce').sum(axis=1)
                    # 选择产品列、金额列和数量列
                    select_cols = [product_col] + in_cols_no_qty + qty_cols_in
                    top5_in = df_with_products.nlargest(5, 'total_in')[select_cols]
                    
                    # 使用实际列名创建表格（产品 + 金额列 + 数量列），不显示 Category
                    table_cols = [product_col] + in_cols_no_qty + qty_cols_in
                    table = doc.add_table(rows=1, cols=len(table_cols))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(table_cols):
                        hdr_cells[i].text = str(col)
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                    
                    # 设置列宽：第一列 50%，其余列平均分配剩余 50%
                    table.columns[0].width = Inches(3.0)  # 第一列占 50%
                    remaining_cols = len(table_cols) - 1
                    if remaining_cols > 0:
                        for i in range(1, len(table_cols)):
                            table.columns[i].width = Inches(3.0 / remaining_cols)  # 平均分配剩余宽度
                    
                    for _, row in top5_in.iterrows():
                        row_cells = table.add_row().cells
                        for i, col in enumerate(table_cols):
                            if col in row:
                                row_cells[i].text = f"{row[col]:,.2f}" if isinstance(row[col], (int, float)) else str(row[col])
                            else:
                                row_cells[i].text = 'N/A'
            
            # 3. Inventory Consumption Analysis 部分
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('●   Inventory Consumption Analysis')
            run.bold = True
            
            total_consumption_amount = 0
            total_consumption_qty = 0
            
            if out_cols:
                numeric_df = df[out_cols].apply(pd.to_numeric, errors='coerce')
                total_consumption_amount = numeric_df.sum().sum()
                qty_cols = [col for col in out_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if qty_cols:
                    numeric_qty_df = df[qty_cols].apply(pd.to_numeric, errors='coerce')
                    total_consumption_qty = numeric_qty_df.sum().sum()
                else:
                    total_consumption_qty = len(df[df[out_cols[0]] != 0])
            
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Total consumption amount for the month: ')
            run = para.add_run(f'${total_consumption_amount:,.2f}')
            run.bold = True
            run = para.add_run(', this covers a total of ')
            run = para.add_run(f'{total_consumption_qty:,.0f}')
            run.bold = True
            run = para.add_run(' items, with the top 5 items by consumption value as follows:')
            
            if product_col and out_cols:
                out_cols_no_qty = [col for col in out_cols if not ("数量" in str(col) or "qty" in str(col).lower())]
                qty_cols_out = [col for col in out_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if out_cols_no_qty:
                    df_with_products = df.copy()
                    df_with_products['total_out'] = df_with_products[out_cols_no_qty].apply(pd.to_numeric, errors='coerce').sum(axis=1)
                    # 选择产品列、金额列和数量列
                    select_cols = [product_col] + out_cols_no_qty + qty_cols_out
                    top5_out = df_with_products.nlargest(5, 'total_out')[select_cols]
                    
                    # 使用实际列名创建表格（产品 + 金额列 + 数量列），不显示 Category
                    table_cols = [product_col] + out_cols_no_qty + qty_cols_out
                    table = doc.add_table(rows=1, cols=len(table_cols))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(table_cols):
                        hdr_cells[i].text = str(col)
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                    
                    # 设置列宽：第一列 50%，其余列平均分配剩余 50%
                    table.columns[0].width = Inches(3.0)  # 第一列占 50%
                    remaining_cols = len(table_cols) - 1
                    if remaining_cols > 0:
                        for i in range(1, len(table_cols)):
                            table.columns[i].width = Inches(3.0 / remaining_cols)  # 平均分配剩余宽度
                    
                    for _, row in top5_out.iterrows():
                        row_cells = table.add_row().cells
                        for i, col in enumerate(table_cols):
                            if col in row:
                                row_cells[i].text = f"{row[col]:,.2f}" if isinstance(row[col], (int, float)) else str(row[col])
                            else:
                                row_cells[i].text = 'N/A'
            
            # 4. Month-End Inventory Balance 部分
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('●   Month-End Inventory Balance')
            run.bold = True
            
            total_ending_amount = 0
            total_ending_qty = 0
            
            if ending_cols:
                numeric_df = df[ending_cols].apply(pd.to_numeric, errors='coerce')
                total_ending_amount = numeric_df.sum().sum()
                qty_cols = [col for col in ending_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if qty_cols:
                    numeric_qty_df = df[qty_cols].apply(pd.to_numeric, errors='coerce')
                    total_ending_qty = numeric_qty_df.sum().sum()
                else:
                    total_ending_qty = len(df[df[ending_cols[0]] != 0])
            
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(f'As of the end of the month, the total inventory value at {site_name} Site is   ')
            run = para.add_run(f'${total_ending_amount:,.2f}')
            run.bold = True
            run = para.add_run(' , covering   ')
            run = para.add_run(f'{total_ending_qty:,.0f}')
            run.bold = True
            run = para.add_run(' items.')
            
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('The top 5 items by ending balance value are as follows:')
            
            if product_col and ending_cols:
                ending_cols_no_qty = [col for col in ending_cols if not ("数量" in str(col) or "qty" in str(col).lower())]
                qty_cols_ending = [col for col in ending_cols if "数量" in str(col) or "qty" in str(col).lower()]
                if ending_cols_no_qty:
                    df_with_products = df.copy()
                    df_with_products['total_ending'] = df_with_products[ending_cols_no_qty].apply(pd.to_numeric, errors='coerce').sum(axis=1)
                    # 选择产品列、金额列和数量列
                    select_cols = [product_col] + ending_cols_no_qty + qty_cols_ending
                    top5_ending = df_with_products.nlargest(5, 'total_ending')[select_cols]
                    
                    # 使用实际列名创建表格（产品 + 金额列 + 数量列），不显示 Category
                    table_cols = [product_col] + ending_cols_no_qty + qty_cols_ending
                    table = doc.add_table(rows=1, cols=len(table_cols))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(table_cols):
                        hdr_cells[i].text = str(col)
                        hdr_cells[i].paragraphs[0].runs[0].bold = True
                    
                    # 设置列宽：第一列 50%，其余列平均分配剩余 50%
                    table.columns[0].width = Inches(3.0)  # 第一列占 50%
                    remaining_cols = len(table_cols) - 1
                    if remaining_cols > 0:
                        for i in range(1, len(table_cols)):
                            table.columns[i].width = Inches(3.0 / remaining_cols)  # 平均分配剩余宽度
                    
                    for _, row in top5_ending.iterrows():
                        row_cells = table.add_row().cells
                        for i, col in enumerate(table_cols):
                            if col in row:
                                row_cells[i].text = f"{row[col]:,.2f}" if isinstance(row[col], (int, float)) else str(row[col])
                            else:
                                row_cells[i].text = 'N/A'
            
            # 5. Please kindly confirm that 部分
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run('Please kindly confirm that:')
            run.bold = True
            
            confirm_items = [
                'All procurement receipts have been recorded accurately.',
                'All consumption transactions have been properly registered.',
                'The month-end inventory quantities are consistent with inventory counts.'
            ]
            
            for item in confirm_items:
                doc.add_paragraph(item, style='List Bullet')
            
            # 保存 Word 文档
            export_dir = os.path.join(os.path.dirname(self.file_path), "导出") if self.file_path else os.path.join(os.getcwd(), "导出")
            if not os.path.exists(export_dir):
                os.makedirs(export_dir)
            
            for file in os.listdir(export_dir):
                if file.endswith(".docx") and "-stock-report_" in file:
                    file_path = os.path.join(export_dir, file)
                    try:
                        os.remove(file_path)
                    except:
                        pass
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            export_file = os.path.join(export_dir, f"{site_name}-stock-report_{timestamp}.docx")
            
            doc.save(export_file)
            
            self.status_label.config(text="Word 报告导出完成", fg="green")
            messagebox.showinfo("成功", f"Word 报告已导出到:\n{export_file}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{str(e)}")
            self.status_label.config(text="导出失败", fg="red")
    
    def show_overview(self):
        if self.df is None:
            messagebox.showwarning("警告", "请先读取 Excel 文件")
            return
        
        try:
            overview_window = tk.Toplevel(self.root)
            overview_window.title("整体概览")
            overview_window.geometry("1000x700")
            
            main_frame = tk.Frame(overview_window)
            main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            category_frame = tk.LabelFrame(main_frame, text="按类别汇总", font=('Arial', 10, 'bold'))
            category_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, ipady=80)
            
            category_tree = ttk.Treeview(category_frame, show="tree headings")
            category_scroll_y = ttk.Scrollbar(category_frame, orient=tk.VERTICAL, command=category_tree.yview)
            category_scroll_x = ttk.Scrollbar(category_frame, orient=tk.HORIZONTAL, command=category_tree.xview)
            category_tree.configure(yscrollcommand=category_scroll_y.set, xscrollcommand=category_scroll_x.set)
            
            category_columns = ["期初", "采购", "耗用", "结存"]
            category_tree["columns"] = category_columns
            category_tree.column("#0", width=200, minwidth=200)
            for col in category_columns:
                category_tree.column(col, width=100, minwidth=100, anchor=tk.E)
                category_tree.heading(col, text=col, command=lambda c=col: self.sort_overview_column(category_tree, c, False))
            
            category_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            category_scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
            category_scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
            
            df = self.filtered_df if self.filtered_df is not None else self.df
            
            category_col = None
            for col in df.columns:
                col_str = str(col)
                if "类别" in col_str or "category" in col_str.lower():
                    category_col = col
                    break
            
            if category_col is None:
                category_col = df.columns[0]
            
            opening_cols = [col for col in df.columns if "opening" in str(col).lower() or "期初" in str(col)]
            in_cols = [col for col in df.columns if ("in" in str(col).lower() or "入库" in str(col)) and not ("opening" in str(col).lower() or "期初" in str(col) or "ending" in str(col).lower() or "期末" in str(col))]
            out_cols = [col for col in df.columns if "out" in str(col).lower() or "出库" in str(col)]
            ending_cols = [col for col in df.columns if "ending" in str(col).lower() or "期末" in str(col)]
            
            categories = sorted(df[category_col].unique())
            category_data = []
            
            for category in categories:
                category_df = df[df[category_col] == category]
                
                opening_amount = 0
                if opening_cols:
                    numeric_df = category_df[opening_cols].apply(pd.to_numeric, errors='coerce')
                    opening_amount = numeric_df.sum().sum()
                
                in_amount = 0
                if in_cols:
                    numeric_df = category_df[in_cols].apply(pd.to_numeric, errors='coerce')
                    in_amount = numeric_df.sum().sum()
                
                out_amount = 0
                if out_cols:
                    numeric_df = category_df[out_cols].apply(pd.to_numeric, errors='coerce')
                    out_amount = numeric_df.sum().sum()
                
                ending_amount = 0
                if ending_cols:
                    numeric_df = category_df[ending_cols].apply(pd.to_numeric, errors='coerce')
                    ending_amount = numeric_df.sum().sum()
                
                category_data.append({
                    "category": category,
                    "opening": opening_amount,
                    "in": in_amount,
                    "out": out_amount,
                    "ending": ending_amount
                })
                
                category_tree.insert("", tk.END, values=[
                    f"{opening_amount:,.2f}",
                    f"{in_amount:,.2f}",
                    f"{out_amount:,.2f}",
                    f"{ending_amount:,.2f}"
                ], text=str(category))
            
            total_opening = sum(item["opening"] for item in category_data)
            total_in = sum(item["in"] for item in category_data)
            total_out = sum(item["out"] for item in category_data)
            total_ending = sum(item["ending"] for item in category_data)
            
            total_item = category_tree.insert("", tk.END, values=[
                f"{total_opening:,.2f}",
                f"{total_in:,.2f}",
                f"{total_out:,.2f}",
                f"{total_ending:,.2f}"
            ], text="合计")
            category_tree.item(total_item, tags=('total',))
            category_tree.tag_configure('total', background='#E0E0E0', font=('Arial', 9, 'bold'))
            
            site_frame = tk.LabelFrame(main_frame, text="按 Site 汇总", font=('Arial', 10, 'bold'))
            site_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, ipady=10)
            
            site_tree = ttk.Treeview(site_frame, show="tree headings")
            site_scroll_y = ttk.Scrollbar(site_frame, orient=tk.VERTICAL, command=site_tree.yview)
            site_scroll_x = ttk.Scrollbar(site_frame, orient=tk.HORIZONTAL, command=site_tree.xview)
            site_tree.configure(yscrollcommand=site_scroll_y.set, xscrollcommand=site_scroll_x.set)
            
            site_columns = ["期初", "采购", "耗用", "结存"]
            site_tree["columns"] = site_columns
            site_tree.column("#0", width=150, minwidth=150)
            for col in site_columns:
                site_tree.column(col, width=100, minwidth=100, anchor=tk.E)
                site_tree.heading(col, text=col, command=lambda c=col: self.sort_overview_column(site_tree, c, False))
            
            site_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            site_scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
            site_scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
            
            location_col = None
            for col in df.columns:
                col_str = str(col)
                if "仓库" in col_str or "location" in col_str.lower() or "site" in col_str.lower():
                    location_col = col
                    break
            
            if location_col is None:
                location_col = df.columns[0]
            
            main_locations = ["HUO", "SAN", "BOS", "SFO", "LAX"]
            site_data = []
            
            for site in self.location_options:
                if site == "其他":
                    mask = ~df[location_col].astype(str).str.contains('|'.join(main_locations), case=False, na=False)
                    site_df = df[mask]
                else:
                    mask = df[location_col].astype(str).str.contains(site, case=False, na=False)
                    site_df = df[mask]
                
                if len(site_df) == 0:
                    continue
                
                opening_amount = 0
                if opening_cols:
                    numeric_df = site_df[opening_cols].apply(pd.to_numeric, errors='coerce')
                    opening_amount = numeric_df.sum().sum()
                
                in_amount = 0
                if in_cols:
                    numeric_df = site_df[in_cols].apply(pd.to_numeric, errors='coerce')
                    in_amount = numeric_df.sum().sum()
                
                out_amount = 0
                if out_cols:
                    numeric_df = site_df[out_cols].apply(pd.to_numeric, errors='coerce')
                    out_amount = numeric_df.sum().sum()
                
                ending_amount = 0
                if ending_cols:
                    numeric_df = site_df[ending_cols].apply(pd.to_numeric, errors='coerce')
                    ending_amount = numeric_df.sum().sum()
                
                site_data.append({
                    "site": site,
                    "opening": opening_amount,
                    "in": in_amount,
                    "out": out_amount,
                    "ending": ending_amount
                })
                
                site_tree.insert("", tk.END, values=[
                    f"{opening_amount:,.2f}",
                    f"{in_amount:,.2f}",
                    f"{out_amount:,.2f}",
                    f"{ending_amount:,.2f}"
                ], text=site)
            
            total_opening = sum(item["opening"] for item in site_data)
            total_in = sum(item["in"] for item in site_data)
            total_out = sum(item["out"] for item in site_data)
            total_ending = sum(item["ending"] for item in site_data)
            
            total_item = site_tree.insert("", tk.END, values=[
                f"{total_opening:,.2f}",
                f"{total_in:,.2f}",
                f"{total_out:,.2f}",
                f"{total_ending:,.2f}"
            ], text="合计")
            site_tree.item(total_item, tags=('total',))
            site_tree.tag_configure('total', background='#E0E0E0', font=('Arial', 9, 'bold'))
            
        except Exception as e:
            messagebox.showerror("错误", f"显示整体概览失败：{str(e)}")
    
    def sort_overview_column(self, tree, col, reverse):
        items = [(tree.item(item)["values"], tree.item(item)["text"]) for item in tree.get_children('')]
        
        def convert_value(val):
            if isinstance(val, str):
                try:
                    return float(val.replace(',', ''))
                except:
                    return 0
            return val
        
        items.sort(key=lambda x: convert_value(x[0][list(tree["columns"]).index(col)]), reverse=reverse)
        
        for item in tree.get_children(''):
            tree.delete(item)
        
        for idx, (values, text) in enumerate(items):
            tree.insert("", tk.END, values=values, text=text)
        
        for i, column in enumerate(tree["columns"]):
            if column == col:
                if reverse:
                    tree.heading(column, text=f"{col} ↑", command=lambda c=column: self.sort_overview_column(tree, c, False))
                else:
                    tree.heading(column, text=f"{col} ↓", command=lambda c=column: self.sort_overview_column(tree, c, True))
            else:
                tree.heading(column, text=str(column), command=lambda c=column: self.sort_overview_column(tree, c, False))
    
    def read_excel(self, file_path=None):
        if file_path is None:
            file_path = self.path_entry.get()
        
        if not file_path:
            messagebox.showwarning("警告", "请先选择或输入文件路径")
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror("错误", "文件不存在")
            return
        
        try:
            self.status_label.config(text="正在读取Excel文件...", fg="orange")
            self.root.update()
            
            self.df = pd.read_excel(file_path)
            self.file_path = file_path
            self.filtered_df = None
            
            for loc in self.location_options:
                self.location_vars[loc].set(True)
            
            self.info_label.config(text=f"文件: {os.path.basename(file_path)} | 行数: {len(self.df)} | 列数: {len(self.df.columns)}")
            
            self.display_data_in_table()
            
            # 计算并显示金额汇总
            self.calculate_amounts()
            
            self.status_label.config(text="Excel文件读取完成", fg="green")
            
        except Exception as e:
            messagebox.showerror("错误", f"读取文件失败: {str(e)}")
            self.status_label.config(text="读取失败", fg="red")
    



def main():
    root = tk.Tk()
    app = StockAnalysisApp(root)
    root.state('zoomed')
    root.mainloop()


if __name__ == "__main__":
    main()
